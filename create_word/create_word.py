from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches,Pt
import time


class DocumentOperation:

    def createNewDocument(self):
        return Document()

    def modifyDocument(self, modify_docu):
        return Document(modify_docu)

    def addTitle(self, docu, title_info, title_size=0):
        return docu.add_heading(text=title_info, level=title_size)
        
    def modifyTitleSize(self, docu, title_info, title_size):
        return docu.add_heading(text=title_info, level=title_size)

    def addTable(self, docu, row, col):
        return docu.add_table(row, col)
        
    def insertPicture(self, docu, pic_path, pic_width:float):
        if pic_width:
            return docu.add_picture(pic_path, width=Inches(pic_width))
        return docu.add_picture(pic_path)

    def addParagraph(self, docu, paragraph_info):
        return docu.add_paragraph(paragraph_info)

    def paragraphIndentation(self, paragraph, indentation_size:float):
        paragraph_format = paragraph.paragraph_format
        return paragraph_format.left_indent == Inches(float(indentation_size))
 
    
    def run(self):
        judge = input("Whether to create a new document?(y/n) ")
        if judge.lower() == "y" or judge == "" or judge.lower() == "yes":
            document = self.createNewDocument()
            while 1:
                print ("""
                        1. add title
                        2. add table
                        3. add paragraph
                        4. add picture
                        5. quit
                       """)
                s = input("Please select the type you want to add:")
                if s:
                    s = int(s)
                    if s == 5:
                        print ("Program exit!")
                        exit(0)
                    elif s == 1:
                        title_info = input("Please add a title to the article:")
                        judge1 = input("Do you need to change the title format? default=0 (y/n)")
                        if judge1.lower() == "y" or judge1 == "" or judge1.lower() == "yes":
                            title_size = int(input("Please enter a title format(0-9):"))
                            self.addTitle(document, title_info, title_size)
                        else:
                            self.addTitle(document, title_info)
                        print ("The title has been added!\n")
                    elif s == 2:
                        judge2 = input("Do you need to add a table?(y/n)")
                        if judge2.lower() == "y" or judge2 == "" or judge2.lower() == "yes":
                            row = int(input("Enter the number of rows you need:"))
                            col = int(input("Enter the number of cols you need:"))
                            self.addTable(document, row, col)
                            print ("The Table has been added!\n")
                    elif s == 3:
                        judge3 = input("Do you want to add a paragraph?(y/n)")
                        if judge3.lower() == "y" or judge3 == "" or judge3.lower() == "yes":
                            paragraph_info = input("Please enter the paragraph content:\n")
                            paragraph = self.addParagraph(document, paragraph_info)
                            paragraph_format = paragraph.paragraph_format
                            while 1:
                                judge6 = input("Do you need custom indentation?(y/n)")
                                if judge6.lower() == "y" or judge6 == "" or judge6.lower() == "yes":
                                    print ("""
                                        1.Entire paragraph indentation
                                        2.First line indentation
                                        3.Right side indent
                                        4.Exit custom indentation
                                        """)
                                    s2 = input("Please choose:")
                                    s2 = int(s2)
                                    if s2 == 1:
                                        value = float(input("Please enter an entire paragraph indentation value:"))
                                        paragraph_format.left_indent = Inches(value)
                                        print ("Entire paragraph indentation {}".format(value))
                                        break
                                    elif s2 == 2:
                                        value = float(input("Please enter an first line indentation value:")) 
                                        paragraph_format.first_line_indent = Inches(value) 
                                        print ("First line indentation {}".format(value))
                                        break
                                    elif s2 == 3:
                                        value = int(input("Please enter an right side indentation value:"))
                                        paragraph_format.right_indent = Pt(value)
                                        print ("Right side indent {}".format(value))
                                        break
                                    elif s2 == 4:
                                        print ("Exit custom indentation")
                                        continue
                                elif judge6.lower() == "n" or judge6.lower() == "no":
                                    print ("Default first line indent!")
                                    paragraph_format = paragraph.paragraph_format
                                    paragraph_format.first_line_indent = Inches(0.3)
                                    break
                                else:
                                    print ("Unrecognized your input, please re-enter...")
                                    continue
                        else:
                            print ("Please re-select...")
                            continue
                    elif s == 4:    
                        judge4 = input("Do you want to insert a picture?(y/n)")
                        if judge4.lower() == "y" or judge4 == "" or judge4.lower() == "yes":
                            pic_path = input("Please enter the picture path:")
                            judge5 = input("Do you need to resize the picture?(y/n)")
                            if judge5.lower() == "y" or judge5 == "" or judge5.lower() == "yes":
                                pic_width = float(input("Please enter the width you want to adjust(float):"))
                                document.add_picture(pic_path, width=Inches(pic_width))
                                print ("Picture width is {}".format(pic_width))
                            else:
                                document.add_picture(pic_path)
                                print ("Keep the default size of the picture!")
                            judge7 = input("Do you need a picture centered?(y/n)")
                            if judge7.lower() == "y" or judge7 == "" or judge7.lower() == "yes":
                                last_paragraph = document.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                print ("The picture is already centered.")
                            else:
                                print ("The picture is left aligned by default.")
                        else:
                            continue
                    else:
                        print("Input is invalid, please re-enter...")
                        continue
                else:
                    print("No characters are entered, please re-enter...")
                    continue
                document.save("/tmp/test_{}.docx".format(time.strftime('%Y-%m-%d',time.localtime(time.time()))))
        else:
            modify_ducu_path = input("Please enter the file path:")
            document = self.modifyDocument(modify_ducu_path)
            judge8 = input("Do you want to read the document content?(y/n)")
            if judge8.lower() == "y" or judge8 == "" or judge8.lower() == "yes":
                for para in document.paragraphs:
                    print(para.text)



if __name__ == "__main__":
    creatdoc = DocumentOperation()
    creatdoc.run()

